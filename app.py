import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from fpdf import FPDF

def generar_informe(datos, incluir_imagen, imagen_path=None):
    doc = Document()
    
    # Establecer formato del documento
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    # Título centrado
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("INFORME TÉCNICO")
    run.bold = True
    run.font.size = Pt(14)
    
    # Datos generales
    doc.add_paragraph(f"Fecha: {datos['fecha']}").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph(f"Destinatario: {datos['destinatario']}").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph(f"Vía: {datos['via']}").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph(f"Remitente: {datos['remitente']}").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph(f"Referencia: {datos['referencia']}").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Secciones del informe
    doc.add_heading('INTRODUCCIÓN', level=2)
    doc.add_paragraph(datos['introduccion']).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_heading('OBJETIVOS', level=2)
    doc.add_paragraph(datos['objetivos']).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_heading('ACTIVIDADES REALIZADAS', level=2)
    for actividad in datos['actividades']:
        doc.add_paragraph(f"- {actividad}").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Tabla de detalles del avance
    doc.add_heading('DETALLES DEL AVANCE', level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Código'
    hdr_cells[1].text = 'Descripción'
    hdr_cells[2].text = 'Estado'
    
    for detalle in datos['detalles_avance']:
        row_cells = table.add_row().cells
        row_cells[0].text = detalle['codigo']
        row_cells[1].text = detalle['descripcion']
        row_cells[2].text = detalle['estado']
    
    # Observaciones
    doc.add_heading('OBSERVACIONES', level=2)
    doc.add_paragraph(datos['observaciones']).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Firmas
    doc.add_heading('FIRMAS', level=2)
    doc.add_paragraph("Firma del Supervisor de Obra: ________________").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph("VoBo del Socio Clínica: ________________").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Guardar documento
    doc_path = f"Informe_Tecnico_{datos['fecha'].replace('/', '-')}.docx"
    doc.save(doc_path)
    return doc_path

def verificar_contraseña():
    contraseña = st.text_input("Ingrese la contraseña:", type="password")
    if contraseña == "f1-f2-f3":
        return True
    else:
        st.error("Contraseña incorrecta")
        return False

def main():
    st.title("📋 Generador de Informe Técnico")
    
    if not verificar_contraseña():
        return
    
    # Formulario para los datos
    fecha = st.text_input("📅 Fecha (DD/MM/AAAA):")
    destinatario = st.text_input("🎯 Destinatario:")
    via = st.text_input("📩 Vía:")
    remitente = st.text_input("✉️ Remitente:")
    referencia = st.text_input("📌 Referencia:")
    introduccion = st.text_area("📖 Introducción:")
    objetivos = st.text_area("🎯 Objetivos:")
    actividades = st.text_area("🛠 Actividades Realizadas (separadas por comas):").split(',')
    observaciones = st.text_area("⚠ Observaciones:")
    
    detalles_avance = []
    num_detalles = st.number_input("📊 Número de Detalles del Avance:", min_value=1, step=1)
    for i in range(int(num_detalles)):
        st.subheader(f"Detalle {i+1}")
        codigo = st.text_input(f"Código {i+1}:")
        descripcion = st.text_input(f"Descripción {i+1}:")
        estado = st.text_input(f"Estado {i+1}:")
        detalles_avance.append({"codigo": codigo, "descripcion": descripcion, "estado": estado})
    
    if st.button("📄 Generar Informe"):
        datos = {
            "fecha": fecha,
            "destinatario": destinatario,
            "via": via,
            "remitente": remitente,
            "referencia": referencia,
            "introduccion": introduccion,
            "objetivos": objetivos,
            "actividades": actividades,
            "detalles_avance": detalles_avance,
            "observaciones": observaciones
        }
        
        doc_path = generar_informe(datos, False, None)
        st.success(f"Informe generado: {doc_path}")
        st.download_button("⬇️ Descargar Informe en Word", open(doc_path, "rb"), doc_path)
        
if __name__ == "__main__":
    main()
